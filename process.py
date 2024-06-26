import matplotlib.pyplot as plt
import pandas as pd
import plotly.graph_objs as go
import numpy as np
import os
import math
import seaborn as sns
import re
from decimal import Decimal
from docx import Document 
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Cm

def get_filepaths(dir):
    filepaths = []
    for filename in os.listdir(dir):
        filepath = os.path.join(dir, filename)
        if os.path.isfile(filepath):
            filepaths.append(filepath)
    return filepaths

def gain_batch_name(filepaths):
    batch_name = filepaths[0][15:]
    pattern = r'\d+'
    numbers = re.findall(pattern, batch_name)
    batch_name = filepaths[0][15:]
    batch_name_1=""
    batch_name_1+=batch_name[:batch_name.find(numbers[0])+numbers[0].__len__()]
    batch_name=batch_name[batch_name.find(numbers[0])+numbers[0].__len__():]
    batch_name_1+=batch_name[:batch_name.find(numbers[1])+numbers[1].__len__()]
    batch_name_2=batch_name_1
    batch_name=batch_name[batch_name.find(numbers[1])+numbers[1].__len__():]
    batch_name=batch_name[batch_name.find(numbers[2])+numbers[2].__len__():]
    total_num = sum([int(re.findall(pattern, i[15:])[3]) for i in filepaths])
    batch_name_1+="_"+str(total_num)+batch_name[batch_name.find(numbers[3])+numbers[3].__len__():-4]
    batch_name=batch_name_1
    batch=batch_name[:2]
    return batch,batch_name,batch_name_2,total_num

def import_scale(batch):
    if batch=="14":
        scale = pd.read_csv("scale_file_CH4.csv").T
        scale.index=["response_resistance_x","response_size_x","discrimination_x","selectivity_x"]
    elif batch=="TV":
        scale = pd.read_csv("scale_file.csv").T
        scale.index=["response_resistance_x","response_size_x","discrimination_x","response_stability_x","restoration_x","baseline_deviation_x","high_discrimination_x"
                     ]
    return scale
def concat_data(filepaths,batch):
    if batch=='TV':
        if pd.read_csv("data/"+filepaths[0],sep=r'\s+',names=[str(i) for i in range(2500)]).columns.__len__()==7:
            data_total = pd.DataFrame(columns=[i for i in range(6)])
            for j,i in enumerate(filepaths):
                data = pd.read_csv("data/"+i,sep=r'\s+',names=[str(i) for i in range(2500)])
                data_total=data_total.transpose().dropna(how="all",axis=0)
                data = data_total[data.sum(axis=1) != 0][1:]
                data["no"] = data.index.values.astype(int)
                data["batch"]=filepaths[j][15:-4]
                data_total= pd.concat([data_total,data])
        else:

            data_total = pd.DataFrame(columns=[i for i in range(7)])
            for j,i in enumerate(filepaths):
                data = pd.read_csv("data/"+i,sep=r'\s+',names=[str(i) for i in range(2500)])
                data=data.transpose().dropna(how="all",axis=0)
                data = data[data.sum(axis=1) != 0][1:]
                data["no"] = data.index.values.astype(int)
                data["batch"]=int(j)
                data_total= pd.concat([data_total,data])

    elif batch=="14":
        data_total = pd.DataFrame(columns=[i for i in range(8)])
        data_total.columns = columns=["0","20","10","5","2","2000","no","batch"]
        for j,i in enumerate(filepaths):
                data = pd.read_csv("data/"+i,sep=r'\s+',names=[str(i) for i in range(2500)])
                data=data.transpose().dropna(how="all",axis=0)
                data = data[data.sum(axis=1) != 0][1:]
                data["no"] = data.index.values.astype(int)
                data["batch"]=filepaths[j][15:22]
                data.columns=["0","20","10","5","2","2000","no","batch"]
                data_total= pd.concat([data_total,data])

    return data_total

def process(batch,batch_name,batch_name_2,scale,data_total,total_num,filepaths):
    if batch=="14":
        response_resistance_x=list(scale[scale.index=="response_resistance_x"].T["response_resistance_x"])
        response_size_x=list(scale[scale.index=="response_size_x"].T["response_size_x"])
        discrimination_x=list(scale[scale.index=="discrimination_x"].T["discrimination_x"])
        selectivity_x=list(scale[scale.index=="selectivity_x"].T["selectivity_x"])

        response_resistance_x=[x for x in response_resistance_x if not math.isnan(x)]
        response_size_x=[x for x in response_size_x if not math.isnan(x)]
        discrimination_x=[x for x in discrimination_x if not math.isnan(x)]
        selectivity_x=[x for x in selectivity_x if not math.isnan(x)]
        response_resistance = data_total["10"]
        response_size = data_total["0"]/data_total["10"]
        discrimination = data_total["10"]/data_total["20"]
        selectivity=(data_total["0"]/data_total["2"])-(data_total["0"]/data_total["2000"])

        y_1=[0]+list(pd.cut(response_resistance,bins=response_resistance_x).value_counts(sort=False,normalize=True).values)
        y_2=[0]+list(pd.cut(response_size,bins=response_size_x).value_counts(sort=False,normalize=True).values)
        y_3=[0]+list(pd.cut(discrimination,bins=discrimination_x).value_counts(sort=False,normalize=True).values)
        y_4=[0]+list(pd.cut(selectivity,bins=selectivity_x).value_counts(sort=False,normalize=True)[1:].values)

        x_1=response_resistance_x[:-1]
        x_1.append(5000)
        x_2=response_size_x
        x_3=discrimination_x
        x_4=selectivity_x[1:]
        x_1[-1]=">4800"

        plt.rcParams['font.sans-serif'] = ['Microsoft YaHei'] 
        plt.rcParams["font.size"] =16
        plt.figure(figsize=(20, 6))
        plt.title("响应电阻(R$_{10}$)")
        plt.xlabel("区间")
        plt.ylabel("频率")
        plt.grid( axis='y', linewidth=0.3)
        plt.xticks(fontsize=10)
        plt.plot(x_1, y_1 ,marker='D', markersize=5)
        plt.savefig('响应电阻.png')
        plt.show()

        print("409600:",pd.cut(response_resistance,bins=response_resistance_x).value_counts(sort=False).values[-1])

        plt.figure(figsize=(20, 6))
        plt.title("响应大小(R$_{0}$/R$_{10}$)")
        plt.xlabel("区间")
        plt.ylabel("频率")
        plt.grid( axis='y', linewidth=0.3)
        plt.plot(x_2, y_2, marker='D', markersize=5)
        plt.savefig('响应大小.png')
        plt.show()

        plt.figure(figsize=(20, 6))
        plt.title("区分度(R$_{10}$/R$_{20}$)")
        plt.xlabel("区间")
        plt.ylabel("频率")
        plt.grid( axis='y', linewidth=0.3)
        plt.plot(x_3, y_3, marker='D', markersize=5)
        plt.savefig('区分度.png')
        plt.show()

        plt.figure(figsize=(20, 6))
        plt.title("选择性[(R$_{0}$/R$_{2}$)-(R$_{0}$/R$_{2000}$)]")
        plt.xlabel("区间")
        plt.ylabel("频率")
        plt.grid( axis='y', linewidth=0.3)
        plt.plot(x_4, y_4, marker='D', markersize=5)
        plt.savefig('选择性.png')
        plt.show()

        no_=data_total.index.values.astype(int)
        data_processed = pd.concat(
                [pd.Series(no_,index=response_resistance.index),
                response_resistance,response_size,discrimination,selectivity,
                data_total["batch"]],axis=1)
        columns=["序号","响应电阻",
            "响应大小",
            "区分度",
            "选择性",
            "部分"
            ]

        data_processed.columns=columns

        step_1=pd.merge(
            data_processed[data_processed["响应电阻"]>=10] , data_processed[data_processed["响应电阻"]<=1000],how="inner"
            )

        step_2=step_1[step_1["响应大小"]>5]

        step_3=step_2[step_2["区分度"]>1.2]
        step_4=step_3[step_3["选择性"]>=0]

        defective_products = []
        defective_products.append(pd.merge(
            data_processed[data_processed["响应电阻"]>=10] , data_processed[data_processed["响应电阻"]<=1000],how="inner"
            ).__len__()/data_processed.__len__())
        defective_products.append(data_processed[data_processed["响应大小"]>5].__len__()/data_processed.__len__())
        defective_products.append(data_processed[data_processed["区分度"]>1.2].__len__()/data_processed.__len__())
        defective_products.append(data_processed[data_processed["选择性"]>0].__len__()/data_processed.__len__())

        step_4_=step_4.copy()
        step_4_.index=step_4["序号"]
        step_4_.index.name=''
        merged_ = pd.concat([data_processed,step_4_])
        merged=merged_.drop_duplicates(keep=False).copy()

        defective_products_1 = []
        defective_products_1.append(pd.concat(
            [merged[merged["响应电阻"]<10] , merged[merged["响应电阻"]>1000]]
            ).__len__()/merged.__len__())
        defective_products_1.append((
            merged[merged["响应大小"]<=5]
            ).__len__()/merged.__len__())
        defective_products_1.append((
            merged[merged["区分度"]<=1.2]
            ).__len__()/merged.__len__())
        defective_products_1.append((
            merged[merged["选择性"]<=0]
            ).__len__()/merged.__len__())
        
        defective_products_2 = []
        defective_products_2.append(pd.concat(
            [merged[merged["响应电阻"]<10] , merged[merged["响应电阻"]>1000]]
            ).__len__())
        defective_products_2.append((
            merged[merged["响应大小"]<=5]
            ).__len__())
        defective_products_2.append((
            merged[merged["区分度"]<=1.2]
            ).__len__())
        defective_products_2.append((
            merged[merged["选择性"]<=0]
            ).__len__())
        
        defective_products_x=["不符合响应电阻","不符合响应大小","不符合区分度","不符合选择性"]
        plt.rcParams["font.size"] =20
        plt.rcParams['font.sans-serif'] = ['Microsoft YaHei'] 
        fig=plt.figure(figsize=(20, 6),dpi=80)
        axis_1 = fig.add_axes((0.1, 0.1, 0.8, 0.8))
        plt.yticks([i for i in np.linspace(0,1,num=6)],[str(round(i,1)*100)+"%"for i in np.linspace(0, 1, num=6)])
        axis_2 = axis_1.twinx()
        axis_1.plot(defective_products_x, defective_products_1, color="#6AA84F", marker="o", linewidth=2)
        axis_1.set_ylabel("比例", size=20) 

        axis_2.bar(defective_products_x,defective_products_2,width=0.5, alpha=0.8,color="#5470C6")
        axis_2.set_ylabel("数量", size=20) 

        plt.grid(True, linewidth=0.3)
        for i in range(defective_products_1.__len__()):  
            plt.text(defective_products_x[i], defective_products_2[i], str(round(Decimal(defective_products_1[i]),2)*100)+"%", ha='center', va= 'bottom')

        plt.title(batch_name_2+"不良品统计情况", size=25)
        plt.legend()
        plt.savefig("不良率.png")
        plt.show()

        filtered_no=[]
        sc= list(step_4["部分"].values)
        for j,i in enumerate(step_4["序号"]):
                if int(i)%64 == 0:
                    filtered_no.append(str(int(i)//64-1)+"+"+str(64)+"+"+str(int(i)%64)+" "+sc[j])
                else:  filtered_no.append(str(int(i)//64)+"+"+str(int(i)%64)+" "+sc[j])

        step_4["序号"]=filtered_no
        step_4.to_excel("processed_CH4_total_"+batch_name+".xlsx")

        document =Document()
        document.styles['Normal'].font.name='楷体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        run=document.add_heading('',level=0).add_run('分析报告')
        run.font.name='微软雅黑'
        _title = document.styles['Title']
        _title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph = document.add_paragraph('本次测试日期为'+filepaths[0][:filepaths[0].find("14")]) 
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph = document.add_paragraph('本次测试批次为：'+batch_name+'，一共测试'+str(total_num)+"个器件")

        paragraph = document.add_paragraph('各参数筛选合格率') 
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER 
        table = document.add_table(rows=1, cols=3, style='Table Grid') 
        table.style.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '指标名称'
        hdr_cells[1].text = '值(良品率)'
        hdr_cells[2].text = '值(不良率)'
        table.cell(0,0).width=Cm(5) 
        table.cell(0,1).width=Cm(10) 

        mlst =["响应电阻",
            "响应大小",
            "区分度",
            "选择性",
            ]

        for i,_row in enumerate(mlst):
            row_cells = table.add_row().cells # 添一行表格元素
            row_cells[0].text = _row
            row_cells[1].text = ' '+str(round(Decimal(defective_products[i]),2)*100)[:5]+"%"
            row_cells[2].text = ' '+str(round(Decimal(1-defective_products[i]),2)*100)[:5]+"%"
            p = row_cells[1].paragraphs[0]
            p.paragraph_format.alignment =WD_ALIGN_PARAGRAPH.JUSTIFY # 单元格文字两端对齐

        paragraph = document.add_paragraph(' ')

        paragraph = document.add_paragraph('参数筛选区间表') 
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER 
        table = document.add_table(rows=1, cols=2, style='Table Grid') 
        table.style.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '指标名称'
        hdr_cells[1].text = '区间范围'
        table.cell(0,0).width=Cm(5) 
        table.cell(0,1).width=Cm(10) 

        mlst =["响应电阻",
            "响应大小",
            "区分度",
            "选择性",
            ]
        sc=["10-1000",">5",">1.2",">0",]

        for i,_row in enumerate(mlst):
            row_cells = table.add_row().cells 
            row_cells[0].text = _row
            row_cells[1].text = ' '+sc[i]
            p = row_cells[1].paragraphs[0]
            p.paragraph_format.alignment =WD_ALIGN_PARAGRAPH.JUSTIFY 

        paragraph = document.add_paragraph(' ')

        paragraph = document.add_paragraph('其中符合4参数器件'+str(step_4.__len__())+"个器件"+"，良品率(占测试总数量)为"+str(round(Decimal(step_4.__len__()/data_total.__len__()),4)*100)[:5]+"%")
        paragraph = document.add_paragraph('其中符合4参数器件:'+str(list(step_4["序号"])))


        document.add_picture('响应电阻.png', width=Cm(15))
        document.add_picture('响应大小.png', width=Cm(15))
        document.add_picture('区分度.png', width=Cm(15))
        document.add_picture('选择性.png', width=Cm(15))
        document.add_picture('不良率.png', width=Cm(15))

        document.save('report/分析报告_'+batch_name+'.docx')

    elif batch=="TV":
        response_resistance_x=list(scale[scale.index=="response_resistance_x"].T["response_resistance_x"])
        response_size_x=list(scale[scale.index=="response_size_x"].T["response_size_x"])
        discrimination_x=list(scale[scale.index=="discrimination_x"].T["discrimination_x"])
        response_stability_x=list(scale[scale.index=="response_stability_x"].T["response_stability_x"])
        restoration_x=list(scale[scale.index=="restoration_x"].T["restoration_x"])
        baseline_deviation_x=list(scale[scale.index=="baseline_deviation_x"].T["baseline_deviation_x"])
        high_discrimination_x=list(scale[scale.index=="high_discrimination_x"].T["high_discrimination_x"])

        response_resistance_x=[x for x in response_resistance_x if not math.isnan(x)]
        response_size_x=[x for x in response_size_x if not math.isnan(x)]
        discrimination_x=[x for x in discrimination_x if not math.isnan(x)]
        response_stability_x=[x for x in response_stability_x if not math.isnan(x)]
        restoration_x=[x for x in restoration_x if not math.isnan(x)]
        baseline_deviation_x=[x for x in baseline_deviation_x if not math.isnan(x)]
        high_discrimination_x=[x for x in high_discrimination_x if not math.isnan(x)]

        if data_total.columns.__len__()==10:
            response_size = data_total[0]/data_total[4]
            discrimination = data_total[4]/data_total[3]
            response_stability=data_total[5]/data_total[4]
            restoration=data_total[6]/data_total[0]
            baseline_deviation=(data_total[0]-data_total[7])/data_total[0]
            high_discrimination=data_total[2]/data_total[1]
            
        if data_total.columns.__len__()==9:
            response_resistance = data_total[2]
            response_size = data_total[0]/data_total[2]
            discrimination = data_total[2]/data_total[1]
            response_stability = data_total[3]/data_total[2]
            restoration = data_total[4]/data_total[0]
            baseline_deviation=(data_total[0]-data_total[6])/data_total[0]

        y_1=[0]+list(pd.cut(response_resistance,bins=response_resistance_x).value_counts(sort=False,normalize=True)[:11].values)
        y_2=[0]+list(pd.cut(response_size,bins=response_size_x).value_counts(sort=False,normalize=True)[:16].values)
        y_3=[0]+list(pd.cut(discrimination,bins=discrimination_x).value_counts(sort=False,normalize=True)[5:25].values)
        y_4=[0]+list(pd.cut(response_stability,bins=response_stability_x).value_counts(sort=False,normalize=True)[10:30].values)
        y_5=[0]+list(pd.cut(restoration,bins=restoration_x).value_counts(sort=False,normalize=True)[17:-1].values)
        y_6=[0]+list(pd.cut(baseline_deviation,bins=baseline_deviation_x).value_counts(sort=False,normalize=True)[15:-1].values)

        x_1=response_resistance_x[:12]
        x_2=response_size_x[:17]
        x_3=discrimination_x[5:26]
        x_4=response_stability_x[10:31]
        x_5=restoration_x[17:-1]
        x_6=baseline_deviation_x[15:-1]
        if data_total.columns.__len__()==8:
            x_7=high_discrimination_x[5:27]
            y_7=[0]+list(pd.cut(high_discrimination,bins=high_discrimination_x).value_counts(sort=False,normalize=True)[5:26].values)

        plt.rcParams['font.sans-serif'] = ['Microsoft YaHei'] 
        plt.rcParams["font.size"] =35
        plt.figure(figsize=(20,12))
        plt.title("响应电阻(R$^{1}$$_{50}$)")
        plt.xlabel("响应电阻值（KΩ）")
        plt.ylabel("频率")
        plt.grid( axis='y', linewidth=0.3)
        plt.plot(x_1, y_1 ,marker='D', markersize=5)
        plt.savefig('响应电阻.png')
        plt.show()

        print("409600:",pd.cut(response_resistance,bins=response_resistance_x).value_counts(sort=False).values[-1])

        plt.figure(figsize=(20, 12))
        plt.title("响应大小(R$^{1}$$_{0}$/R$^{1}$$_{50}$)")
        plt.xlabel("响应大小")
        plt.ylabel("频率")
        plt.grid( axis='y', linewidth=0.3)
        plt.plot(x_2, y_2, marker='D', markersize=5)
        plt.savefig('响应大小.png')
        plt.show()

        plt.figure(figsize=(20, 12))
        plt.title("区分度(R$^{1}$$_{50}$/R$_{80}$)")
        plt.xlabel("区分度")
        plt.ylabel("频率")
        plt.grid( axis='y', linewidth=0.3)
        plt.plot(x_3, y_3, marker='D', markersize=5)
        plt.savefig('区分度.png')
        plt.show()

        plt.figure(figsize=(20, 12))
        plt.title("响应稳定性(R$^{2}$$_{50}$/R$^{1}$$_{50}$)")
        plt.xlabel("响应稳定性")
        plt.ylabel("频率")
        plt.grid( axis='y', linewidth=0.3)
        plt.plot(x_4, y_4, marker='D', markersize=5)
        plt.savefig('响应稳定性.png')
        plt.show()

        plt.figure(figsize=(20, 12))
        plt.title("恢复程度(R$_{53}$/R$^{1}$$_{0}$)")
        plt.xlabel("恢复程度")
        plt.ylabel("频率")
        plt.grid( axis='y', linewidth=0.3)
        plt.plot(x_5, y_5, marker='D', markersize=5)
        plt.savefig('恢复程度.png')
        plt.show()

        plt.figure(figsize=(20, 6))
        plt.title("基线偏差[(R$^{1}$$_{0}$-R$^{2}$$_{0}$)/R$^{1}$$_{0}$]")
        plt.xlabel("基线偏差")
        plt.ylabel("频率")
        plt.grid( axis='y', linewidth=0.3)
        plt.plot(x_6, y_6, marker='D', markersize=5)
        plt.savefig('基线偏差.png')
        plt.show()

        if data_total.columns.__len__()==8:
            plt.title("大浓度区分度值(R$_{200}$/R$_{100}$)")
            plt.xlabel("大浓度区分度值")
            plt.ylabel("频率")
            plt.grid( axis='y', linewidth=0.3)
            plt.plot(x_7, y_7, marker='D', markersize=5)
            plt.savefig('大浓度区分度值.png')
            plt.show()

        plt.close()

        no_=data_total.index.values.astype(int)
        if data_total.columns.__len__()==9:
            data_processed = pd.concat(
                [pd.Series(no_,index=response_resistance.index),
                response_resistance,response_size,discrimination,response_stability,restoration,baseline_deviation
                ,data_total["batch"]],axis=1)
            columns=["序号","响应电阻",
            "响应大小",
            "区分度",
            "响应稳定性",
            "恢复程度",
            "基线偏差",
            "部分"
            ]
        if data_total.columns.__len__()==10:
            data_processed = pd.concat(
                [pd.Series(no_,index=response_resistance.index),
                response_resistance,response_size,discrimination,response_stability,restoration,baseline_deviation,high_discrimination
                ,data_total["batch"]],axis=1)
            columns=["序号","响应电阻",
            "响应大小",
            "区分度",
            "响应稳定性",
            "恢复程度",
            "基线偏差",
            "大浓度区分度",
            "部分"
            ]
        data_processed.columns=columns
        step_1=pd.merge(
            data_processed[data_processed["响应电阻"]>=3] , data_processed[data_processed["响应电阻"]<=30],how="inner"
            )

        step_2=pd.merge(
            step_1[step_1["响应大小"]>=2] , step_1[step_1["响应大小"]<=15],how="inner"
            )


        step_3=pd.merge(
            step_2[step_2["区分度"]>=1.1] , step_2[step_2["区分度"]<=1.5],how="inner"
            )


        step_4=pd.merge(
            step_3[step_3["响应稳定性"]>=0.95] , step_3[step_3["响应稳定性"]<=1.05],how="inner"
            )


        step_5=step_4[step_4["恢复程度"]>0.75]

        step_6=pd.merge(
            step_5[step_5["基线偏差"]>=-0.2] , step_5[step_5["基线偏差"]<=0.1],how="inner"
            )


        if data_total.columns.__len__()==10:
            step_7=step_6[step_6["大浓度区分度"]>1.3]

        step_6_=step_6.copy()
        step_6_.index=step_6["序号"]
        step_6_.index.name=''
        merged = pd.concat([data_processed,step_6_])
        merged.drop_duplicates(keep=False,inplace=True)

        if data_total.columns.__len__()==10:
            step_7_=step_7.copy()
            step_7_.index=step_7["序号"]
            step_7_.index.name=''
            merged = pd.concat([data_processed,step_7_])
            merged.drop_duplicates(keep=False,inplace=True)

        defective_products_1 = []
        defective_products_1.append(pd.concat(
            [merged[merged["响应电阻"]<3] , merged[merged["响应电阻"]>30]]
            ).__len__()/merged.__len__())
        defective_products_1.append(pd.concat(
            [merged[merged["响应大小"]<2] , merged[merged["响应大小"]>15]]
            ).__len__()/merged.__len__())
        defective_products_1.append(pd.concat(
            [merged[merged["区分度"]<1.1] , merged[merged["区分度"]>1.5]]
            ).__len__()/merged.__len__())
        defective_products_1.append(pd.concat(
            [merged[merged["响应稳定性"]<0.95] , merged[merged["响应稳定性"]>1.05]]
            ).__len__()/merged.__len__())
        defective_products_1.append(merged[merged["恢复程度"]<=0.75].__len__()/merged.__len__())
        defective_products_1.append(pd.concat(
            [merged[merged["基线偏差"]<-0.2] , merged[merged["基线偏差"]>0.1]]
            ).__len__()/merged.__len__())
        if data_total.columns.__len__()==10:
            defective_products_1.append(merged[merged["大浓度区分度"]<=1.3].__len__()/merged.__len__())

        defective_products_2 = []
        defective_products_2.append(pd.concat(
            [merged[merged["响应电阻"]<3] , merged[merged["响应电阻"]>30]]
            ).__len__())
        defective_products_2.append(pd.concat(
            [merged[merged["响应大小"]<2] , merged[merged["响应大小"]>15]]
            ).__len__())
        defective_products_2.append(pd.concat(
            [merged[merged["区分度"]<1.1] , merged[merged["区分度"]>1.5]]
            ).__len__())
        defective_products_2.append(pd.concat(
            [merged[merged["响应稳定性"]<0.95] , merged[merged["响应稳定性"]>1.05]]
            ).__len__())
        defective_products_2.append(merged[merged["恢复程度"]<=0.75].__len__())
        defective_products_2.append(pd.concat(
            [merged[merged["基线偏差"]<-0.2] , merged[merged["基线偏差"]>0.1]]
            ).__len__())
        if data_total.columns.__len__()==10:
            defective_products_2.append(merged[merged["大浓度区分度"]<=1.3].__len__())
        defective_products_x=["不符合响应电阻","不符合响应大小","不符合区分度","不符合响应稳定性","不符合恢复程度","不符基线偏差"]
        if data_total.columns.__len__()==10:
            defective_products_x=["不符合响应电阻","不符合响应大小","不符合区分度","不符合响应稳定性","不符合恢复程度","不符基线偏差","不符合大浓度区分度"]
        plt.rcParams["font.size"] =27

        plt.rcParams['font.sans-serif'] = ['SimSun']
        fig=plt.figure(figsize=(22, 12),dpi=120)
        axis_1 = fig.add_axes((0.1, 0.1, 0.8, 0.8))
        plt.yticks([i for i in np.linspace(0,1,num=6)],[str(round(i,1)*100)[:3]+"%"for i in np.linspace(0, 1, num=6)])

        axis_2 = axis_1.twinx()
        axis_1.plot(defective_products_x, defective_products_1, color="#6AA84F", marker="o", linewidth=4.5)
        axis_1.set_ylabel("比例", size=30)
        axis_1.xaxis.set_tick_params(pad=20) 

        colors = ["#5470C6" for i in range(defective_products_x.__len__())]
        colors[:3]=["#A60000" for i in range(3)]
        axis_2.bar(defective_products_x,defective_products_2,width=0.5, alpha=0.8,color=colors)
        axis_2.set_ylabel("数量", size=30) 
        plt.ylim(0,1800)
        plt.grid(True, linewidth=0.3)
        for i in range(defective_products_1.__len__()):  
            plt.text(defective_products_x[i], defective_products_2[i]+50, str(round(Decimal(defective_products_1[i]),4)*100)[:4]+"%", ha='center', va= 'bottom',fontweight='bold')

        plt.title(batch_name_2+"不良品统计情况", size=45,pad=40)
        plt.legend()
        plt.savefig("不良率.png")
        plt.show()

        filtered_no=[]
        for j,i in enumerate(step_6["序号"]):
                if int(i)%64 == 0:
                    filtered_no.append(str(int(i)//64-1)+"+"+str(64)+"_"+str(int(step_6["部分"][j])+1)+"部分")
                else:  filtered_no.append(str(int(i)//64)+"+"+str(int(i)%64)+"_"+str(int(step_6["部分"][j])+1)+"部分")
        step_6["序号"]=filtered_no

        if data_total.columns.__len__()==10:
            filtered_no=[]
            for j,i in enumerate(step_7["序号"]):
                if int(i)%64 == 0:
                    filtered_no.append(str(int(i)//64-1)+"+"+str(64)+"_"+str(int(step_7["部分"][j])+1)+"部分")
                else:  filtered_no.append(str(int(i)//64)+"+"+str(int(i)%64)+"_"+str(int(step_7["部分"][j])+1)+"部分")
            step_7["序号"]=filtered_no
        step_6.to_excel("processed_total_TVOC.xlsx")
        if data_total.columns.__len__()==10:
            step_7.to_excel("processed_total_TVOC.xlsx")

        defective_products = []
        defective_products.append(pd.merge(
            data_processed[data_processed["响应电阻"]>=3] , data_processed[data_processed["响应电阻"]<=30],how="inner"
            ).__len__()/data_processed.__len__())
        defective_products.append(pd.merge(
            data_processed[data_processed["响应大小"]>=2] , data_processed[data_processed["响应大小"]<=15],how="inner"
            ).__len__()/data_processed.__len__())
        defective_products.append(pd.merge(
            data_processed[data_processed["区分度"]>=1.1] , data_processed[data_processed["区分度"]<=1.5],how="inner"
            ).__len__()/data_processed.__len__())
        defective_products.append(pd.merge(
            data_processed[data_processed["响应稳定性"]>=0.95] , data_processed[data_processed["响应稳定性"]<=1.05],how="inner"
            ).__len__()/data_processed.__len__())
        defective_products.append(data_processed[data_processed["恢复程度"]>0.75].__len__()/data_processed.__len__())
        defective_products.append(pd.merge(
            data_processed[data_processed["基线偏差"]>=-0.2] , data_processed[data_processed["基线偏差"]<=0.1],how="inner"
            ).__len__()/data_processed.__len__())
        if data_total.columns.__len__()==10:
            defective_products.append(data_processed[data_processed["大浓度区分度"]>1.3].__len__()/data_processed.__len__())

        document =Document()
        document.styles['Normal'].font.name='楷体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        run=document.add_heading('',level=0).add_run('分析报告')
        run.font.name='微软雅黑'
        _title = document.styles['Title']
        _title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph = document.add_paragraph('本次测试日期为'+filepaths[0][:5]) 
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph = document.add_paragraph('本次测试批次为：'+batch_name+'，一共测试'+str(total_num)+"个器件")

        paragraph = document.add_paragraph('各参数筛选合格率') 
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER 
        table = document.add_table(rows=1, cols=2, style='Table Grid') 
        table.style.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '指标名称'
        hdr_cells[1].text = '值'
        table.cell(0,0).width=Cm(5) 
        table.cell(0,1).width=Cm(10) 

        if data_total.columns.__len__()==10:
            mlst =["序号","响应电阻",
            "响应大小",
            "区分度",
            "响应稳定性",
            "恢复程度",
            "基线偏差",
            "大浓度区分度",
            ]
        if data_total.columns.__len__()==9:
            mlst =["响应电阻",
            "响应大小",
            "区分度",
            "响应稳定性",
            "恢复程度",
            "基线偏差",
            ]
        for i,_row in enumerate(mlst):
            row_cells = table.add_row().cells # 添一行表格元素
            row_cells[0].text = _row
            row_cells[1].text = ' '+str(round(defective_products[i],2)*100)+"%"
            p = row_cells[1].paragraphs[0]
            p.paragraph_format.alignment =WD_ALIGN_PARAGRAPH.JUSTIFY # 单元格文字两端对齐

        paragraph = document.add_paragraph(' ')

        paragraph = document.add_paragraph('参数筛选区间表') 
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER 
        table = document.add_table(rows=1, cols=2, style='Table Grid') 
        table.style.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '指标名称'
        hdr_cells[1].text = '区间范围'
        table.cell(0,0).width=Cm(5) 
        table.cell(0,1).width=Cm(10) 

        if data_total.columns.__len__()==10:
            mlst =["序号","响应电阻",
            "响应大小",
            "区分度",
            "响应稳定性",
            "恢复程度",
            "基线偏差",
            "大浓度区分度",
            ]
            sc=["3-30K","2-15","1.1-1.5","0.95~1.05","大于0.75","-0.2~0.1","大于1.3"]
        if data_total.columns.__len__()==9:
            mlst =["响应电阻",
            "响应大小",
            "区分度",
            "响应稳定性",
            "恢复程度",
            "基线偏差",
            ]
            sc=["3-30K","2-15","1.1-1.5","0.95~1.05","大于0.75","-0.2~0.1"]
        for i,_row in enumerate(mlst):
            row_cells = table.add_row().cells 
            row_cells[0].text = _row
            row_cells[1].text = ' '+sc[i]
            p = row_cells[1].paragraphs[0]
            p.paragraph_format.alignment =WD_ALIGN_PARAGRAPH.JUSTIFY 

        paragraph = document.add_paragraph(' ')

        paragraph = document.add_paragraph('其中符合6参数器件'+str(step_6.__len__())+"个器件"+"，良品率(占测试总数量)为"+str(round(step_6.__len__()/data_total.__len__(),4)*100)+"%")
        paragraph = document.add_paragraph('其中符合6参数器件:'+str(list(step_6["序号"])))

        if data_total.columns.__len__()==10:
            paragraph = document.add_paragraph('其中符合7参数器件'+str(step_7.__len__())+"个器件"+"，良品率(占测试总数量)为"+str(round(step_7.__len__()/data_total.__len__(),4)*100)+"%")
            paragraph = document.add_paragraph('其中符合7参数器件:'+str(list(step_7["序号"])))



        document.add_picture('响应电阻.png', width=Cm(15))
        document.add_picture('响应大小.png', width=Cm(15))
        document.add_picture('响应稳定性.png', width=Cm(15))
        document.add_picture('区分度.png', width=Cm(15))
        document.add_picture('恢复程度.png', width=Cm(15))
        document.add_picture('基线偏差.png', width=Cm(15))

        if data_total.columns.__len__()==10:
            document.add_picture('大浓度区分度值.png', width=Cm(15))

        document.add_picture('不良率.png', width=Cm(15))

        document.save('report/分析报告_'+batch_name+'.docx')

def import_data_one(batch,filename):
    if batch=="TV":
        path_="data/"+filename
        data = pd.read_csv(path_,sep=r'\s+',names=[str(i) for i in range(2500)])
        data=data.transpose().dropna(how="all",axis=0)
        data = data[data.sum(axis=1) != 0][1:]
        batch_name=path_[20:-4]
        return batch_name,data
    elif batch=="14":
        path_="data/"+filename
        data = pd.read_csv(path_,sep=r'\s+',names=[str(i) for i in range(2500)])
        data=data.transpose().dropna(how="all",axis=0)
        data = data[data.sum(axis=1) != 0][1:]
        data.columns=["0","20","10","5","2","2000"]
        batch_name=path_[20:-4]
        return batch_name,data

def process_one_by_one(batch,batch_name,scale,data,filename):
    if batch=="TV":
        if data.columns.__len__()==8:
            response_resistance = data[4]
            response_size = data[0]/data[4]
            discrimination = data[4]/data[3]
            response_stability=data[5]/data[4]
            restoration=data[6]/data[0]
            baseline_deviation=(data[0]-data[7])/data[0]
            high_discrimination=data[2]/data[1]
        
        if data.columns.__len__()==7:
            response_resistance = data[2]
            response_size = data[0]/data[2]
            discrimination = data[2]/data[1]
            response_stability = data[3]/data[2]
            restoration = data[4]/data[0]
            baseline_deviation=(data[0]-data[6])/data[0]

        response_resistance_x=list(scale[scale.index=="response_resistance_x"].T["response_resistance_x"])
        response_size_x=list(scale[scale.index=="response_size_x"].T["response_size_x"])
        discrimination_x=list(scale[scale.index=="discrimination_x"].T["discrimination_x"])
        response_stability_x=list(scale[scale.index=="response_stability_x"].T["response_stability_x"])
        restoration_x=list(scale[scale.index=="restoration_x"].T["restoration_x"])
        baseline_deviation_x=list(scale[scale.index=="baseline_deviation_x"].T["baseline_deviation_x"])
        high_discrimination_x=list(scale[scale.index=="high_discrimination_x"].T["high_discrimination_x"])

        response_resistance_x=[x for x in response_resistance_x if not math.isnan(x)]
        response_size_x=[x for x in response_size_x if not math.isnan(x)]
        discrimination_x=[x for x in discrimination_x if not math.isnan(x)]
        response_stability_x=[x for x in response_stability_x if not math.isnan(x)]
        restoration_x=[x for x in restoration_x if not math.isnan(x)]
        baseline_deviation_x=[x for x in baseline_deviation_x if not math.isnan(x)]
        high_discrimination_x=[x for x in high_discrimination_x if not math.isnan(x)]

        y_1=[0]+list(pd.cut(response_resistance,bins=response_resistance_x).value_counts(sort=False,normalize=True)[:11].values)
        y_2=[0]+list(pd.cut(response_size,bins=response_size_x).value_counts(sort=False,normalize=True)[:16].values)
        y_3=[0]+list(pd.cut(discrimination,bins=discrimination_x).value_counts(sort=False,normalize=True)[5:25].values)
        y_4=[0]+list(pd.cut(response_stability,bins=response_stability_x).value_counts(sort=False,normalize=True)[10:30].values)
        y_5=[0]+list(pd.cut(restoration,bins=restoration_x).value_counts(sort=False,normalize=True)[17:-1].values)
        y_6=[0]+list(pd.cut(baseline_deviation,bins=baseline_deviation_x).value_counts(sort=False,normalize=True)[15:-1].values)

        x_1=response_resistance_x[:12]
        x_2=response_size_x[:17]
        x_3=discrimination_x[5:26]
        x_4=response_stability_x[10:31]
        x_5=restoration_x[17:-1]
        x_6=baseline_deviation_x[15:-1]
        if data.columns.__len__()==8:
            x_7=high_discrimination_x[5:27]
            y_7=[0]+list(pd.cut(high_discrimination,bins=high_discrimination_x).value_counts(sort=False,normalize=True)[5:26].values)

        plt.rcParams['font.sans-serif'] = ['Microsoft YaHei'] 
        plt.rcParams["font.size"] =16
        plt.figure(figsize=(20, 12))
        plt.title("响应电阻(R$^{1}$$_{50}$)")
        plt.xlabel("响应电阻(KΩ)")
        plt.ylabel("频率")
        plt.grid( axis='y', linewidth=0.3)
        plt.plot(x_1, y_1 ,marker='D', markersize=5)
        plt.savefig('响应电阻.png')
  

        print("409600:",pd.cut(response_resistance,bins=response_resistance_x).value_counts(sort=False).values[-1])

        plt.figure(figsize=(20, 12))
        plt.title("响应大小(R$^{1}$$_{0}$/R$^{1}$$_{50}$)")
        plt.xlabel("响应大小")
        plt.ylabel("频率")
        plt.grid( axis='y', linewidth=0.3)
        plt.plot(x_2, y_2, marker='D', markersize=5)
        plt.savefig('响应大小.png')
     

        plt.figure(figsize=(20, 12))
        plt.title("区分度(R$^{1}$$_{50}$/R$_{80}$)")
        plt.xlabel("区分度")
        plt.ylabel("频率")
        plt.grid( axis='y', linewidth=0.3)
        plt.plot(x_3, y_3, marker='D', markersize=5)
        plt.savefig('区分度.png')
   

        plt.figure(figsize=(20, 12))
        plt.title("响应稳定性(R$^{2}$$_{50}$/R$^{1}$$_{50}$)")
        plt.xlabel("响应稳定性")
        plt.ylabel("频率")
        plt.grid( axis='y', linewidth=0.3)
        plt.plot(x_4, y_4, marker='D', markersize=5)
        plt.savefig('响应稳定性.png')


        plt.figure(figsize=(20, 12))
        plt.title("恢复程度(R$_{53}$/R$^{1}$$_{0}$)")
        plt.xlabel("恢复程度")
        plt.ylabel("频率")
        plt.grid( axis='y', linewidth=0.3)
        plt.plot(x_5, y_5, marker='D', markersize=5)
        plt.savefig('恢复程度.png')
  

        plt.figure(figsize=(20, 12))
        plt.title("基线偏差[(R$^{1}$$_{0}$-R$^{2}$$_{0}$)/R$^{1}$$_{0}$]")
        plt.xlabel("基线偏差")
        plt.ylabel("频率")
        plt.grid( axis='y', linewidth=0.3)
        plt.plot(x_6, y_6, marker='D', markersize=5)
        plt.savefig('基线偏差.png')


        if data.columns.__len__()==8:
            plt.figure(figsize=(20, 12))
            plt.title("大浓度区分度值(R$_{200}$/R$_{100}$)")
            plt.xlabel("大浓度区分度值")
            plt.ylabel("频率")
            plt.grid( axis='y', linewidth=0.3)
            plt.plot(x_7, y_7, marker='D', markersize=5)
            plt.savefig('大浓度区分度值.png')
  

        no_=data.index.values.astype(int)
        if data.columns.__len__()==7:
            data_processed = pd.concat(
                [pd.Series(no_,index=response_resistance.index),
                response_resistance,response_size,discrimination,response_stability,restoration,baseline_deviation
                ],axis=1)
            columns=["序号","响应电阻",
            "响应大小",
            "区分度",
            "响应稳定性",
            "恢复程度",
            "基线偏差",
            ]
        if data.columns.__len__()==8:
            data_processed = pd.concat(
                [pd.Series(no_,index=response_resistance.index),
                response_resistance,response_size,discrimination,response_stability,restoration,baseline_deviation,high_discrimination
                ],axis=1)
            columns=["序号","响应电阻",
            "响应大小",
            "区分度",
            "响应稳定性",
            "恢复程度",
            "基线偏差",
            "大浓度区分度"
            ]
        data_processed.columns=columns

        step_1=pd.merge(
            data_processed[data_processed["响应电阻"]>=3] , data_processed[data_processed["响应电阻"]<=30],how="inner"
            )

        step_2=pd.merge(
            step_1[step_1["响应大小"]>=2] , step_1[step_1["响应大小"]<=15],how="inner"
            )

        step_3=pd.merge(
            step_2[step_2["区分度"]>=1.1] , step_2[step_2["区分度"]<=1.5],how="inner"
            )

        step_4=pd.merge(
            step_3[step_3["响应稳定性"]>=0.95] , step_3[step_3["响应稳定性"]<=1.05],how="inner"
            )
   
        step_5=step_4[step_4["恢复程度"]>0.75]

        step_6=pd.merge(
            step_5[step_5["基线偏差"]>=-0.2] , step_5[step_5["基线偏差"]<=0.1],how="inner"
            )

        if data.columns.__len__()==8:
            step_7=step_6[step_6["大浓度区分度"]>1.3]

        step_6_=step_6.copy()
        step_6_.index=step_6["序号"]
        step_6_.index.name=''
        merged = pd.concat([data_processed,step_6_])
        merged.drop_duplicates(keep=False,inplace=True)

        if data.columns.__len__()==8:
            step_7_=step_7.copy()
            step_7_.index=step_7["序号"]
            step_7_.index.name=''
            merged = pd.concat([data_processed,step_7_])
            merged.drop_duplicates(keep=False,inplace=True)

        defective_products_1 = []
        defective_products_1.append(pd.concat(
            [merged[merged["响应电阻"]<3] , merged[merged["响应电阻"]>30]]
            ).__len__()/merged.__len__())
        defective_products_1.append(pd.concat(
            [merged[merged["响应大小"]<2] , merged[merged["响应大小"]>15]]
            ).__len__()/merged.__len__())
        defective_products_1.append(pd.concat(
            [merged[merged["区分度"]<1.1] , merged[merged["区分度"]>1.5]]
            ).__len__()/merged.__len__())
        defective_products_1.append(pd.concat(
            [merged[merged["响应稳定性"]<0.95] , merged[merged["响应稳定性"]>1.05]]
            ).__len__()/merged.__len__())
        defective_products_1.append(merged[merged["恢复程度"]<=0.75].__len__()/merged.__len__())
        defective_products_1.append(pd.concat(
            [merged[merged["基线偏差"]<-0.2] , merged[merged["基线偏差"]>0.1]]
            ).__len__()/merged.__len__())
        if data.columns.__len__()==8:
            defective_products_1.append(merged[merged["大浓度区分度"]<=1.3].__len__()/merged.__len__())

        defective_products_2 = []
        defective_products_2.append(pd.concat(
            [merged[merged["响应电阻"]<3] , merged[merged["响应电阻"]>30]]
            ).__len__())
        defective_products_2.append(pd.concat(
            [merged[merged["响应大小"]<2] , merged[merged["响应大小"]>15]]
            ).__len__())
        defective_products_2.append(pd.concat(
            [merged[merged["区分度"]<1.1] , merged[merged["区分度"]>1.5]]
            ).__len__())
        defective_products_2.append(pd.concat(
            [merged[merged["响应稳定性"]<0.95] , merged[merged["响应稳定性"]>1.05]]
            ).__len__())
        defective_products_2.append(merged[merged["恢复程度"]<=0.75].__len__())
        defective_products_2.append(pd.concat(
            [merged[merged["基线偏差"]<-0.2] , merged[merged["基线偏差"]>0.1]]
            ).__len__())
        if data.columns.__len__()==8:
            defective_products_2.append(merged[merged["大浓度区分度"]<=1.3].__len__())

        defective_products_x=["不符合响应电阻","不符合响应大小","不符合区分度","不符合响应稳定性","不符合恢复程度","不符基线偏差"]
        if data.columns.__len__()==8:
            defective_products_x=["不符合响应电阻","不符合响应大小","不符合区分度","不符合响应稳定性","不符合恢复程度","不符基线偏差","不符合大浓度区分度"]


        plt.rcParams["font.size"] =27
        plt.rcParams['font.sans-serif'] = ['SimSun']
        fig=plt.figure(figsize=(22, 12),dpi=120)
        axis_1 = fig.add_axes((0.1, 0.1, 0.8, 0.8))
        plt.yticks([i for i in np.linspace(0,1,num=6)],[str(round(i,1)*100)[:3]+"%"for i in np.linspace(0, 1, num=6)])

        axis_2 = axis_1.twinx()
        axis_1.plot(defective_products_x, defective_products_1, color="#6AA84F", marker="o", linewidth=4.5)
        axis_1.set_ylabel("比例", size=30)
        axis_1.xaxis.set_tick_params(pad=20) 

        colors = ["#5470C6" for i in range(defective_products_x.__len__())]
        colors[:3]=["#A60000" for i in range(3)]
        axis_2.bar(defective_products_x,defective_products_2,width=0.5, alpha=0.8,color=colors)
        axis_2.set_ylabel("数量", size=30) 
        plt.ylim(0,896)
        plt.grid(True, linewidth=0.3)
        for i in range(defective_products_1.__len__()):  
            plt.text(defective_products_x[i], defective_products_2[i]+50, str(round(Decimal(defective_products_1[i]),4)*100)[:4]+"%", ha='center', va= 'bottom',fontweight='bold')

        plt.title(batch_name+"不良品统计情况", size=45,pad=40)
        plt.legend()
        plt.savefig("不良率.png")

        filtered_no=[]
        step_6 = step_6.sort_values(by="序号")
        for i in step_6["序号"]:
                if int(i)%64 == 0:
                    filtered_no.append(str(int(i)//64-1)+"+"+str(64))
                else:  filtered_no.append(str(int(i)//64)+"+"+str(int(i)%64))
        step_6["序号"]=filtered_no

        if data.columns.__len__()==8:
            filtered_no=[]
            for i in step_7["序号"]:
                if int(i)%64 == 0:
                    filtered_no.append(str(int(i)//64-1)+"+"+str(64))
                else:  filtered_no.append(str(int(i)//64)+"+"+str(int(i)%64))
            step_7["序号"]=filtered_no
        step_6.to_excel("processed_TVOC.xlsx")
        if data.columns.__len__()==8:
            step_7.to_excel("processed_TVOC.xlsx")

        defective_products = []
        defective_products.append(pd.merge(
            data_processed[data_processed["响应电阻"]>=3] , data_processed[data_processed["响应电阻"]<=30],how="inner"
            ).__len__()/data_processed.__len__())
        defective_products.append(pd.merge(
            data_processed[data_processed["响应大小"]>=2] , data_processed[data_processed["响应大小"]<=15],how="inner"
            ).__len__()/data_processed.__len__())
        defective_products.append(pd.merge(
            data_processed[data_processed["区分度"]>=1.1] , data_processed[data_processed["区分度"]<=1.5],how="inner"
            ).__len__()/data_processed.__len__())
        defective_products.append(pd.merge(
            data_processed[data_processed["响应稳定性"]>=0.95] , data_processed[data_processed["响应稳定性"]<=1.05],how="inner"
            ).__len__()/data_processed.__len__())
        defective_products.append(data_processed[data_processed["恢复程度"]>0.75].__len__()/data_processed.__len__())
        defective_products.append(pd.merge(
            data_processed[data_processed["基线偏差"]>=-0.2] , data_processed[data_processed["基线偏差"]<=0.1],how="inner"
            ).__len__()/data_processed.__len__())
        
        document =Document()
        document.styles['Normal'].font.name='楷体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        run=document.add_heading('',level=0).add_run('分析报告')
        run.font.name='微软雅黑'
        _title = document.styles['Title']
        _title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph = document.add_paragraph('本次测试日期为'+filename[5:filename.find("T")]) 
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph = document.add_paragraph('本次测试批次为'+batch_name+'，一共测试'+str(data.__len__())+"个器件")

        paragraph = document.add_paragraph('各参数筛选合格率') 
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER 
        table = document.add_table(rows=1, cols=3, style='Table Grid') 
        table.style.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '指标名称'
        hdr_cells[1].text = '值(良品率)'
        hdr_cells[2].text = '值(不良率)'
        table.cell(0,0).width=Cm(5) 
        table.cell(0,1).width=Cm(10) 

        if data.columns.__len__()==8:
            mlst =["序号","响应电阻",
            "响应大小",
            "区分度",
            "响应稳定性",
            "恢复程度",
            "基线偏差",
            "大浓度区分度",
            ]
        if data.columns.__len__()==7:
            mlst =["响应电阻",
            "响应大小",
            "区分度",
            "响应稳定性",
            "恢复程度",
            "基线偏差",
            ]
        for i,_row in enumerate(mlst):
            row_cells = table.add_row().cells # 添一行表格元素
            row_cells[0].text = _row
            row_cells[1].text = ' '+str(round(defective_products[i],2)*100)+"%"
            row_cells[2].text = ' '+str(round(Decimal(1-defective_products[i]),2)*100)[:5]+"%"
            p = row_cells[1].paragraphs[0]
            p.paragraph_format.alignment =WD_ALIGN_PARAGRAPH.JUSTIFY # 单元格文字两端对齐

        paragraph = document.add_paragraph(' ')

        paragraph = document.add_paragraph('参数筛选区间表') 
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER 
        table = document.add_table(rows=1, cols=2, style='Table Grid') 
        table.style.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '指标名称'
        hdr_cells[1].text = '区间范围'
        table.cell(0,0).width=Cm(5) 
        table.cell(0,1).width=Cm(10) 

        if data.columns.__len__()==8:
            mlst =["序号","响应电阻",
            "响应大小",
            "区分度",
            "响应稳定性",
            "恢复程度",
            "基线偏差",
            "大浓度区分度",
            ]
            sc=["3-30","2-15","1.1-1.5","0.95~1.05","大于0.75","-0.2~0.1","大于1.3"]
        if data.columns.__len__()==7:
            mlst =["响应电阻",
            "响应大小",
            "区分度",
            "响应稳定性",
            "恢复程度",
            "基线偏差",
            ]
            sc=["3-30","2-15","1.1-1.5","0.95~1.05","大于0.75","-0.2~0.1"]
        for i,_row in enumerate(mlst):
            row_cells = table.add_row().cells 
            row_cells[0].text = _row
            row_cells[1].text = ' '+sc[i]
            p = row_cells[1].paragraphs[0]
            p.paragraph_format.alignment =WD_ALIGN_PARAGRAPH.JUSTIFY 

        paragraph = document.add_paragraph(' ')

        paragraph = document.add_paragraph('其中符合6参数器件'+str(step_6.__len__())+"个器件"+"，良品率(占测试总数量)为"+str(round(step_6.__len__()/data.__len__(),5)*100)+"%")
        paragraph = document.add_paragraph('其中符合6参数器件:'+str(list(step_6["序号"])))

        if data.columns.__len__()==8:
            paragraph = document.add_paragraph('其中符合7参数器件'+str(step_7.__len__())+"个器件"+"，良品率(占测试总数量)为"+str(round(step_7.__len__()/data.__len__(),5)*100)+"%")
            paragraph = document.add_paragraph('其中符合7参数器件:'+str(list(step_7["序号"])))
        document.add_picture('响应电阻.png', width=Cm(15))
        document.add_picture('响应大小.png', width=Cm(15))
        document.add_picture('区分度.png', width=Cm(15))
        document.add_picture('恢复程度.png', width=Cm(15))
        document.add_picture('基线偏差.png', width=Cm(15))
        document.add_picture('不良率.png', width=Cm(15))

        if data.columns.__len__()==8:
            document.add_picture('大浓度区分度值.png', width=Cm(15))

        document.save('report/分析报告_'+batch_name+'.docx')


    elif batch=="14":
        response_resistance_x=list(scale[scale.index=="response_resistance_x"].T["response_resistance_x"])
        response_size_x=list(scale[scale.index=="response_size_x"].T["response_size_x"])
        discrimination_x=list(scale[scale.index=="discrimination_x"].T["discrimination_x"])
        selectivity_x=list(scale[scale.index=="selectivity_x"].T["selectivity_x"])

        response_resistance_x=[x for x in response_resistance_x if not math.isnan(x)]
        response_size_x=[x for x in response_size_x if not math.isnan(x)]
        discrimination_x=[x for x in discrimination_x if not math.isnan(x)]
        selectivity_x=[x for x in selectivity_x if not math.isnan(x)]

        response_resistance = data["10"]
        response_size = data["0"]/data["10"]
        discrimination = data["10"]/data["20"]
        selectivity=(data["0"]/data["2"])-(data["0"]/data["2000"])

        no_=data.index.values.astype(int)

        data_processed = pd.concat(
                [pd.Series(no_,index=response_resistance.index),
                response_resistance,response_size,discrimination,selectivity
                ],axis=1)
        columns=["序号","响应电阻",
            "响应大小",
            "区分度",
            "选择性",
            ]
        data_processed.columns=columns

        y_1=[0]+list(pd.cut(response_resistance,bins=response_resistance_x).value_counts(sort=False,normalize=True).values)
        y_2=[0]+list(pd.cut(response_size,bins=response_size_x).value_counts(sort=False,normalize=True).values)
        y_3=[0]+list(pd.cut(discrimination,bins=discrimination_x).value_counts(sort=False,normalize=True).values)
        y_4=[0]+list(pd.cut(selectivity,bins=selectivity_x).value_counts(sort=False,normalize=True)[1:].values)

        x_1=response_resistance_x[:-1]
        x_1.append(5000)
        x_2=response_size_x
        x_3=discrimination_x
        x_4=selectivity_x[1:]
        x_1[-1]=">4800"

        plt.rcParams['font.sans-serif'] = ['Microsoft YaHei'] 
        plt.rcParams["font.size"] =16
        plt.figure(figsize=(20, 6))
        plt.title("响应电阻(R$_{10}$)")
        plt.xlabel("区间")
        plt.ylabel("频率")
        plt.xticks(fontsize=10)
        plt.grid( axis='y', linewidth=0.3)
        plt.plot(x_1, y_1 ,marker='D', markersize=5)
        plt.savefig('响应电阻.png')

        print("409600:",pd.cut(response_resistance,bins=response_resistance_x).value_counts(sort=False).values[-1])

        plt.figure(figsize=(20, 6))
        plt.title("响应大小(R$_{0}$/R$_{10}$)")
        plt.xlabel("区间")
        plt.ylabel("频率")
        plt.grid( axis='y', linewidth=0.3)
        plt.plot(x_2, y_2, marker='D', markersize=5)
        plt.savefig('响应大小.png')

        plt.figure(figsize=(20, 6))
        plt.title("区分度(R$_{10}$/R$_{20}$)")
        plt.xlabel("区间")
        plt.ylabel("频率")
        plt.grid( axis='y', linewidth=0.3)
        plt.plot(x_3, y_3, marker='D', markersize=5)
        plt.savefig('区分度.png')

        plt.figure(figsize=(20, 6))
        plt.title("选择性[(R$_{0}$/R$_{2}$)-(R$_{0}$/R$_{2000}$)]")
        plt.xlabel("区间")
        plt.ylabel("频率")
        plt.grid( axis='y', linewidth=0.3)
        plt.plot(x_4, y_4, marker='D', markersize=5)
        plt.savefig('选择性.png')

        step_1=pd.merge(
            data_processed[data_processed["响应电阻"]>=10] , data_processed[data_processed["响应电阻"]<=1000],how="inner"
            )

        step_2=step_1[step_1["响应大小"]>5]
        step_3=step_2[step_2["区分度"]>1.2]
        step_4=step_3[step_3["选择性"]>=0]

        filtered_no=[]
        step_4 = step_4.sort_values(by="序号")
        for i in step_4["序号"]:
                if int(i)%64 == 0:
                    filtered_no.append(str(int(i)//64-1)+"+"+str(64))
                else:  filtered_no.append(str(int(i)//64)+"+"+str(int(i)%64))
        step_4["序号"]=filtered_no
        step_4.to_excel("processed_CH4_"+batch_name+".xlsx")

        defective_products = []
        defective_products.append(pd.merge(
            data_processed[data_processed["响应电阻"]>=10] , data_processed[data_processed["响应电阻"]<=1000],how="inner"
            ).__len__()/data_processed.__len__())
        defective_products.append(data_processed[data_processed["响应大小"]>5].__len__()/data_processed.__len__())
        defective_products.append(data_processed[data_processed["区分度"]>1.2].__len__()/data_processed.__len__())
        defective_products.append(data_processed[data_processed["选择性"]>0].__len__()/data_processed.__len__())

        document =Document()
        document.styles['Normal'].font.name='楷体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        run=document.add_heading('',level=0).add_run('分析报告')
        run.font.name='微软雅黑'
        _title = document.styles['Title']
        _title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph = document.add_paragraph('本次测试日期为'+filename[5:filename.find("C")]) 
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph = document.add_paragraph('本次测试批次为'+batch_name+'，一共测试'+str(data.__len__())+"个器件")

        paragraph = document.add_paragraph('各参数筛选合格率') 
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER 
        table = document.add_table(rows=1, cols=3, style='Table Grid') 
        table.style.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '指标名称'
        hdr_cells[1].text = '值(良品率)'
        hdr_cells[2].text = '值(不良率)'
        table.cell(0,0).width=Cm(5) 
        table.cell(0,1).width=Cm(10) 

        mlst =["响应电阻",
            "响应大小",
            "区分度",
            "选择性",
            ]

        for i,_row in enumerate(mlst):
            row_cells = table.add_row().cells # 添一行表格元素
            row_cells[0].text = _row
            row_cells[1].text = ' '+str(round(Decimal(defective_products[i]),2)*100)[:5]+"%"
            row_cells[2].text = ' '+str(round(Decimal(1-defective_products[i]),2)*100)[:5]+"%"
            p = row_cells[1].paragraphs[0]
            p.paragraph_format.alignment =WD_ALIGN_PARAGRAPH.JUSTIFY # 单元格文字两端对齐

        paragraph = document.add_paragraph(' ')

        paragraph = document.add_paragraph('参数筛选区间表') 
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER 
        table = document.add_table(rows=1, cols=2, style='Table Grid') 
        table.style.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '指标名称'
        hdr_cells[1].text = '区间范围'
        table.cell(0,0).width=Cm(5) 
        table.cell(0,1).width=Cm(10) 

        mlst =["响应电阻",
            "响应大小",
            "区分度",
            "选择性",
            ]
        sc=["10-1000",">5",">1.2",">0",]

        for i,_row in enumerate(mlst):
            row_cells = table.add_row().cells 
            row_cells[0].text = _row
            row_cells[1].text = ' '+sc[i]
            p = row_cells[1].paragraphs[0]
            p.paragraph_format.alignment =WD_ALIGN_PARAGRAPH.JUSTIFY 

        paragraph = document.add_paragraph(' ')

        paragraph = document.add_paragraph('其中符合4参数器件'+str(step_4.__len__())+"个器件"+"，良品率(占测试总数量)为"+str(round(Decimal(step_4.__len__()/data.__len__()),4)*100)[:5]+"%")
        paragraph = document.add_paragraph('其中符合4参数器件:'+str(list(step_4["序号"])))

        document.add_picture('响应电阻.png', width=Cm(15))
        document.add_picture('响应大小.png', width=Cm(15))
        document.add_picture('区分度.png', width=Cm(15))
        document.add_picture('选择性.png', width=Cm(15))



        document.save('report/分析报告_'+batch_name+'.docx')



filepaths = get_filepaths("data")
for i,j in enumerate(filepaths):
    filepaths[i]=j[5:]
batch,batch_name,batch_name_2,total_num=gain_batch_name(filepaths)
scale = import_scale(batch)
data_total =concat_data(filepaths,batch)
process(batch,batch_name,batch_name_2,scale,data_total,total_num,filepaths)

for filename in filepaths:
    batch_name,data=import_data_one(batch,filename)
    process_one_by_one(batch,batch_name,scale,data,filename)